import tkinter as tk
from tkinter import simpledialog
from tkinter import messagebox
import openai
import re
import docx
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload
doc = docx.Document()
doc = Document()
def ventana_principal():
    # Función para manejar el evento click del botón 1
    def boton1_click():
        ventana_notas()

    # Función para mostrar la ventana de notas
    def ventana_notas():
        # Función para generar la nota
        def generar_nota():
            # Obtenemos el texto ingresado por el usuario
            texto = texto_nota.get(1.0, tk.END)

            # Generamos el título
            
            titulo = generar_titulo(texto)
            doc.add_paragraph(titulo)

            # Generamos el encabezado
            encabezado = generar_encabezado(texto)
            doc.add_paragraph(encabezado)

            # Generamos la frase clave
            fraseclave = obtener_frase_clave(titulo)
            doc.add_paragraph(fraseclave)
            # Generamos el cuerpo
            cuerpo= generar_cuerpo(titulo,texto,fraseclave)
            # Generamos la metadescripcion
            metadescripcion = generar_metadescripcion(cuerpo)
            doc.add_paragraph(metadescripcion)

            # Generamos la bajada de facebook
            bajada = generar_bajada(cuerpo)
            doc.add_paragraph(bajada)

            # Generamos las etiquetas
            etiquetas = "destacado, " + generar_hashtags(cuerpo)
            doc.add_paragraph(etiquetas)
            
             # Habilitamos la caja de texto del resultado para que pueda ser seleccionado y copiado
           
            resultado.config(state="normal") # habilitamos la caja de texto para que se pueda editar
            resultado.delete(1.0, tk.END) # limpiamos la caja de texto del resultado anterior
            resultado.insert(tk.END, f"Título: {titulo}\nEncabezado:{encabezado}\n{cuerpo}\n\nFrase clave:{fraseclave}\nMetadescripcion:{metadescripcion}\n Bajada:{bajada}\n Etiquetas: {etiquetas}") # mostramos el resultado
            resultado.config(state="disabled") # deshabilitamos la caja de texto para que no se pueda editar
        # Función para generar el título
        def generar_titulo(texto):
            openai.api_key = "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"

            modelo_engine = "text-davinci-002"
            prompt = (f"Genera un título de menos de 8 palabras para este artículo:\n"
                      f"{texto}\n"
                      "Título:")
            completions = openai.Completion.create(engine=modelo_engine, prompt=prompt, max_tokens=32)
            mensaje = completions.choices[0].text.strip()

            # Verificamos si el título tiene más de 5 palabras
            if len(mensaje.split()) > 5:
                respuesta = tk.messagebox.askquestion("Título largo", f"El título tiene {len(mensaje.split())} palabras Titulo:'{mensaje}'. ¿Quieres un título más corto?",master=ventana)
                if respuesta == "yes":
                    mensaje = generar_titulo(texto)
    
            return mensaje
        def generar_encabezado(texto, length=60):
            openai.api_key = "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"
            modelo_engine = "text-davinci-002"
            
            if len(texto.split()) >= 50:
                prompt = (f"Genera un encabezado para este artículo:\n"
                          f"{texto}\n"
                          f"Longitud máxima del encabezado: {length} palabras.\n"
                          "Encabezado:")
            else:
                prompt = (f"Genera un encabezado para este artículo:\n"
                          f"{texto}\n"
                          f"Longitud máxima del encabezado: {length} palabras.\n"
                          "Encabezado:")
                          
            completions = openai.Completion.create(engine=modelo_engine, prompt=prompt, max_tokens=64)
            mensaje = completions.choices[0].text.strip()
            mensaje = re.sub('[^0-9a-zA-ZáéíóúÁÉÍÓÚ \n\.,#]', '', mensaje)
            
            encabezado = mensaje.split("\n")[0]
            
            if len(encabezado.split()) > 30 or len(encabezado.split()) < 10:
                if len(encabezado.split()) > 20:
                    for i in range(3):
                        respuesta = messagebox.askyesno("Pregunta", f"El encabezado '{encabezado}' tiene más de 20 palabras. ¿Deseas uno más corto?", master=ventana)
                        if respuesta:
                            encabezado = generar_encabezado(texto, length=20)
                            break
                else:
                    respuesta = messagebox.askyesno("Pregunta", f"El encabezado '{encabezado}' tiene menos de 10 palabras. ¿Deseas uno más largo?", master=ventana)
                    if respuesta:
                        encabezado = generar_encabezado(texto, length=80)
            
            return encabezado
        def obtener_frase_clave(titulo):
            openai.api_key = "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"
            modelo_engine = "text-davinci-002"
            completado = False
            intentos = 0
            frase_clave = ""

            # Verificar que el título tenga al menos 4 palabras
            palabras_titulo = titulo.split()
            if len(palabras_titulo) < 4:
                return frase_clave

            while not completado and intentos < 10:
               
                # Obtener las sugerencias de frase clave objetivo
                prompt = (f"Por favor dame una frase clave objetivo de menos de 4 palabras que esté contenida en el siguiente título:\n"
                          f"{titulo}\n"
                          f"Frase clave objetivo: ")
                completions = openai.Completion.create(engine=modelo_engine, prompt=prompt, max_tokens=64, n=5)
                opciones = [c.text.strip() for c in completions.choices]
                
                # Validar cada sugerencia
                for opcion in opciones:
                    print("estoy aca ", opcion)
                    # Verificar que la opción tenga menos de 4 palabras
                    palabras_opcion = opcion.split()
                    if len(palabras_opcion) >= 4:
                        continue

                    # Verificar que la opción esté contenida en el título
                    if opcion not in titulo:
                        continue

                    # La opción cumple con los requisitos, usarla como frase clave
                    frase_clave = opcion
                    completado = True
                    break

                intentos += 1

            if completado:
                return frase_clave
            else:
                # Pedir opciones al usuario
                while True:
                    print("Ninguna de las opciones cumple con los requisitos, por favor ingrese una frase clave de menos de 4 palabras que esté contenida en el título:")
                    frase_clave = input().strip()
                    palabras_opcion = frase_clave.split()
                    if len(palabras_opcion) >= 4:
                        continue

                    if frase_clave not in titulo:
                        continue

                    return frase_clave

        # Función para generar el cuerpo de la nota con chatgpt 
        def generar_cuerpo(titulo,texto, fraseclave):
            
            openai.api_key = "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"
            modelo_engine = "text-davinci-002"
            completions = []
            
            # Dividir el texto del usuario en segmentos de 1024 tokens
            texto_segmentado = [texto[i:i+1024] for i in range(0, len(texto), 2000)]
            
            for segmento in texto_segmentado:
                if len(segmento.split()) >= 50:
                    prompt = (f"Genera el cuerpo para este artículo:\n"
                              f"{texto}\n"
                              f"{segmento}\n")
                else:
                    prompt = (f"Genera el cuerpo para este artículo:\n"
                              f"{texto}\n"
                              f"{segmento}\n")
                
                # Generar texto con GPT-3 para el segmento actual
                completions.append(openai.Completion.create(engine=modelo_engine, prompt=prompt, max_tokens=2000))
                
            # Concatenar los resultados de cada segmento
            mensaje = "".join([c.choices[0].text.strip() for c in completions])
            if len(texto) > len(mensaje):
                print("estoy aca1")
                # Mostramos el título y la nota generada
                resultado.config(state="normal") # habilitamos la caja de texto para que se pueda editar
                resultado.delete(1.0, tk.END) # limpiamos la caja de texto del resultado anterior
                resultado.insert(tk.END, f"Título: {titulo}\n\n{mensaje}") # mostramos el resultado
                resultado.config(state="disabled") # deshabilitamos la caja de texto para que no se pueda editar
                
                while True:
                    extenso_cuerpo = mensaje
                    respuesta = messagebox.askyesno("¿Generar cuerpo más extenso?", "¿Quieres que el cuerpo sea más extenso?", master=ventana)
                    if respuesta:
                        
                        
                         # Encontrar la última parte generada del texto original
                        ultima_parte_generada = mensaje.rsplit(segmento)[-1]
                        # Pasar la siguiente parte del texto original a GPT-3
                        siguiente_parte = texto[len(ultima_parte_generada):len(ultima_parte_generada)+1024]
                        prompt = (f"Por favor genera la siguiente parte del cuerpo para este artículo:\n"
                                  f"{siguiente_parte}\n")
                        completions = openai.Completion.create(engine=modelo_engine, prompt=prompt, max_tokens=2000)
                        siguiente_parte_generada = completions.choices[0].text.strip()
                        siguiente_parte_generada = re.sub('[^0-9a-zA-ZáéíóúÁÉÍÓÚ \n\.,#]', '', siguiente_parte_generada)
                        # Concatenar la siguiente parte generada al cuerpo actual
                        mensaje += siguiente_parte_generada
                        if len(extenso_cuerpo) < len(mensaje):
                            extenso_cuerpo = mensaje
                            if len(extenso_cuerpo) > len(texto):
                                
                                break
                        else:
                           resultado.config(state="normal") # habilitamos la caja de texto para que se pueda editar
                           resultado.delete(1.0, tk.END) # limpiamos la caja de texto del resultado anterior
                           resultado.insert(tk.END, f"Título: {titulo}\n\n{extenso_cuerpo}") # mostramos el resultado
                           resultado.config(state="disabled") # deshabilitamos la caja de texto para que no se pueda editar 
                    else:
                        print("estoy aca2")
                        return extenso_cuerpo
                
            return mensaje
        def generar_bajada(cuerpo):
            openai.api_key = "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"
            modelo_engine = "text-davinci-002"
            prompt = (f"Genera una bajada para este contenido de Facebook:\n"
                      f"{cuerpo}\n"
                      "Bajada:")
            bajada_corta = ''
            for _ in range(5):
                completions = openai.Completion.create(engine=modelo_engine, prompt=prompt, max_tokens=300)
                mensaje = completions.choices[0].text.strip()
                mensaje = re.sub('[^0-9a-zA-ZáéíóúÁÉÍÓÚ \n\.,#]', '', mensaje)
                if len(mensaje) <= 300:
                    return mensaje
                elif len(mensaje) < len(bajada_corta) or not bajada_corta:
                    bajada_corta = mensaje
            return bajada_corta

            
        def generar_metadescripcion(cuerpo):
            openai.api_key = "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"
            modelo_engine = "text-davinci-002"
            
            mensaje_mas_corto = None
            for i in range(10):
                if mensaje_mas_corto is not None and len(mensaje_mas_corto) <= 155:
                    break
                
                if len(cuerpo.split()) >= 50:
                    prompt = (f"Genera una metadescripción para esta página:\n"
                              f"{cuerpo}\n"
                              f"Longitud máxima de la metadescripción: 155 caracteres.\n"
                              "Metadescripción:")
                else:
                    prompt = (f"Genera una metadescripción para esta página:\n"
                              f"{cuerpo}\n"
                              f"Longitud máxima de la metadescripción: 155 caracteres.\n"
                              "Metadescripción:")
                              
                completions = openai.Completion.create(engine=modelo_engine, prompt=prompt, max_tokens=100)
                mensaje = completions.choices[0].text.strip()
                mensaje = re.sub('[^0-9a-zA-ZáéíóúÁÉÍÓÚ \n\.,#]', '', mensaje)
                
                if mensaje_mas_corto is None or len(mensaje) < len(mensaje_mas_corto):
                    mensaje_mas_corto = mensaje
            
            return mensaje_mas_corto
        def generar_hashtags(cuerpo):
            openai.api_key = "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"
            modelo_engine = "text-davinci-002"
            prompt = (f"Genera hashtags para este artículo:\n"
                      f"{cuerpo}\n"
                      "Hashtags:")
            completions = openai.Completion.create(engine=modelo_engine, prompt=prompt, max_tokens=200)
            mensaje = completions.choices[0].text.strip()
            mensaje = re.sub('[^0-9a-zA-ZáéíóúÁÉÍÓÚ \n\.,#]', '', mensaje)
            hashtags = [h.strip() for h in mensaje.split("#") if h != '']
            hashtags_str = ", ".join(hashtags)
            return hashtags_str
        # Autenticación de Google Drive
        SCOPES = ['https://www.googleapis.com/auth/drive.file']
        SERVICE_ACCOUNT_FILE = r'C:\Users\joaqu\AppData\Local\Programs\Python\Python310\generar nota\generar-notas-c003aaa97f11.json'


        creds = None
        creds = service_account.Credentials.from_service_account_file(
            SERVICE_ACCOUNT_FILE, scopes=SCOPES)

        # Función para guardar la nota en un archivo de Word y subirla a Google Drive
        def guardar_nota():
            

            # Preguntar al usuario si desea subir el archivo a Google Drive
            respuesta = messagebox.askquestion("Guardar nota en Google Drive", "¿Desea guardar la nota generada en Google Drive?")

            if respuesta == 'yes':
                # Pedir al usuario que ingrese el nombre del archivo para guardar en Google Drive
                nombre_archivo = tk.simpledialog.askstring("Nombre del archivo", "Por favor, ingrese un nombre para el archivo a guardar en Google Drive")
                doc.save(nombre_archivo + '.docx')
                # Subir el archivo a Google Drive
                try:
                    service = build('drive', 'v3', credentials=creds)

                    file_metadata = {'name': nombre_archivo + '.docx', 'parents': ['1BJUMi04_hybrWIfRBe-_pIqq1uAQI2R1']}
                    media = MediaFileUpload(nombre_archivo + '.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
                    messagebox.showinfo("Archivo guardado", f"El archivo '{nombre_archivo}.docx' ha sido guardado en Google Drive. ID: {file.get('id')}")
                except HttpError as error:
                    messagebox.showerror("Error", f"No se pudo subir el archivo a Google Drive. Error: {error}")

        # Creamos una ventana de nivel superior (top-level window)
        ventana = tk.Toplevel(root)
        ventana.title("Generador de notas")
        ventana.geometry("400x400")
    
        # Agregamos una caja de texto para ingresar el texto de la nota
        texto_nota = tk.Text(ventana, height=10)
        texto_nota.pack(pady=10)

        # Agregamos un botón para generar la nota
        boton_generar = tk.Button(ventana, text="Generar nota", command=generar_nota)
        boton_generar.pack(pady=10)
        # Agregar un botón para guardar la nota
        boton_guardar = tk.Button(ventana, text="Guardar nota", command=guardar_nota)
        boton_guardar.pack(pady=10)
        # Agregamos una caja de texto para mostrar el resultado
        resultado = tk.Text(ventana, height=20)
        resultado.pack(pady=10)
        
        # Deshabilitamos la caja de texto del resultado para que no se pueda editar
        resultado.config(state="disabled")

    # Creamos la ventana principal
    root = tk.Tk()
    root.title("Ventana principal")
    root.geometry("300x300")

  # Agregamos el botón 1
    boton1 = tk.Button(root, text="Boton 1", command=boton1_click)
    boton1.pack(pady=10)

  # Agregamos el botón 2
    boton2 = tk.Button(root, text="Boton 2")
    boton2.pack(pady=10)

  # Función para manejar el evento click del botón 2
    def boton2_click():
        print("Click en el botón 2")

  # Asociamos la función boton2_click() al evento click del botón 2
    boton2.config(command=boton2_click)
    
  # Iniciamos el loop de la aplicación
    
    root.mainloop()
ventana_principal()
